import random
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import msoffcrypto
from file_generators.base.base_file import BaseFile
from file_generators.utils.helpers import generate_random_id


class WordFile(BaseFile):
    """Word document generator with malware signature preservation support."""

    def get_file_extension(self) -> str:
        return ".docx"

    def generate(self, data: str, location: str = "center",
                 encrypt: bool = False, preserve_content: bool = False, **kwargs) -> bytes:
        """
        Generate Word document with data.

        Args:
            data: Text content for document
            location: Text alignment ("top-left", "top-right", "center", "justified")
            encrypt: Whether to encrypt the document
            preserve_content: If True, preserves content exactly (critical for EICAR)
            **kwargs: Additional parameters (ignored)

        Returns:
            Word document bytes
        """
        try:
            if preserve_content:
                print("WordFile: Preserving content exactly (EICAR mode)")
                # CRITICAL: For EICAR, use exact content - NO random ID
                full_data = data
                use_random_color = False
            else:
                print(f"WordFile: Standard mode, location={location}")
                # Standard mode: add random ID for tracking
                full_data = data + "\n" + generate_random_id()
                use_random_color = True

            # Create document
            doc = Document()
            paragraph = doc.add_paragraph()

            # Add text with appropriate color
            if use_random_color:
                # Standard mode: random color
                random_color = RGBColor(
                    random.randint(0, 255),
                    random.randint(0, 255),
                    random.randint(0, 255)
                )
            else:
                # Preserved mode: use black color
                random_color = RGBColor(0, 0, 0)  # Black

            run = paragraph.add_run(full_data)
            run.font.color.rgb = random_color

            # Set alignment
            alignments = {
                "top-left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "top-right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "justified": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            paragraph.alignment = alignments.get(location, WD_PARAGRAPH_ALIGNMENT.CENTER)

            # Save to bytes
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            if encrypt:
                print(f"WordFile: Encrypting Word document with password")
                encrypted_bytes = self._encrypt_document(doc_bytes)
                if preserve_content:
                    print(f"WordFile: Generated encrypted Word with preserved content ({len(encrypted_bytes)} bytes)")
                return encrypted_bytes
            else:
                byte_data = doc_bytes.getvalue()
                if preserve_content:
                    print(f"WordFile: Generated Word with preserved content ({len(byte_data)} bytes)")
                else:
                    print(f"WordFile: Generated standard Word with {location} alignment ({len(byte_data)} bytes)")
                return byte_data

        except Exception as e:
            print(f"Error generating Word document: {e}")
            raise

    def _encrypt_document(self, doc_bytes: BytesIO) -> bytes:
        """Encrypt Word document with password."""
        try:
            office_file = msoffcrypto.OfficeFile(doc_bytes)
            encrypted_buffer = BytesIO()
            office_file.encrypt(self.password, encrypted_buffer)
            encrypted_buffer.seek(0)
            return encrypted_buffer.getvalue()
        except Exception as e:
            print(f"Error encrypting Word document: {e}")
            raise
